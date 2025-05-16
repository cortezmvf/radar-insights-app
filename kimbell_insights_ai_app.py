import streamlit as st
import pandas as pd
from google.cloud import bigquery
from docx import Document
from datetime import datetime
import openai
import os
import io
import plotly.graph_objects as go
import plotly.express as px
import time 

# === Setup ===
st.title("📊 Monthly Marketing Analysis")
openai.api_key = st.secrets["OPENAI_API_KEY"]
project_id = "radar-377104"

# Define your Assistant ID globally so both Run Analysis and follow-ups can use it
assistant_id = "asst_lpwUYcZBCm6cUXQsfgWovZKO"

with open("gcp_key.json", "w") as f:
    f.write(st.secrets["GOOGLE_APPLICATION_CREDENTIALS"])
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "gcp_key.json"

system_message = "You are a strategic marketing insights assistant. Follow all formatting and analytical instructions in the user prompt."

# === Session State Setup ===
if "analysis_output" not in st.session_state:
    st.session_state.analysis_output = None
if "followup_count" not in st.session_state:
    st.session_state.followup_count = 0
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = []

# === Month and Metric Selectors ===
month_options = ["2025-01", "2025-02", "2025-03", "2025-04"]
col_a, col_b = st.columns([1, 1])
with col_a:
    selected_month = st.selectbox("Select Month", options=month_options)
with col_b:
    metric_options = ["Impressions", "Clicks", "Sessions", "Revenue"]
    selected_metric = st.selectbox("Select Metric", options=metric_options, key="chart_metric_selector")

# === Load Data from GBQ ===
@st.cache_data(ttl=3600)
def load_data(month):
    client = bigquery.Client(project=project_id)
    query = f"""
        SELECT * 
        FROM `radar-377104.Schaefer_Kimbell.ai_insights_monthly_table`
        WHERE Month_String = '{month}'
    """
    return client.query(query).to_dataframe()

df = load_data(selected_month)

if df.empty:
    st.warning("No data returned for the selected month.")
elif "Campaign_Group" not in df.columns or "Spend" not in df.columns:
    st.error("Required columns ('Campaign_Group', 'Spend') not found in the data.")
else:
    # === Chart Section ===
    metric_by_group = df.groupby("Campaign_Group")[selected_metric].sum().reset_index()
    metric_by_group = metric_by_group[metric_by_group[selected_metric] > 0]

    st.markdown(f"<h3 style='margin-bottom: 0.5rem;'>{selected_metric} by Campaign Group</h3>", unsafe_allow_html=True)

    col1, col2 = st.columns([1, 2])
    with col1:
        fig_pie = go.Figure(data=[go.Pie(
            labels=metric_by_group["Campaign_Group"],
            values=metric_by_group[selected_metric],
            hole=0.5,
            textinfo="percent+label",
            hoverinfo="skip",
            showlegend=False
        )])
        fig_pie.update_layout(height=300, margin=dict(t=0, b=20))
        st.plotly_chart(fig_pie, use_container_width=True)

    with col2:
        fig_bar = px.bar(
            df,
            x="Channel",
            y=selected_metric,
            color="Campaign_Group",
            barmode="group",
            hover_data={selected_metric: ":,.0f"},
            title=" "
        )
        fig_bar.update_layout(showlegend=False, yaxis_tickformat=",", height=300, margin=dict(t=0, b=20))
        st.plotly_chart(fig_bar, use_container_width=True)

    # === Run Analysis Button ===
    if st.session_state.analysis_output is None:
        if st.button("Run Analysis"):
            with st.spinner(f"Running analysis for {selected_month}..."):
                csv_data = df.to_csv(index=False)
                st.session_state.csv_data = csv_data

                # Create thread and store in session
                thread = openai.beta.threads.create()
                st.session_state.analysis_thread_id = thread.id

                openai.beta.threads.messages.create(
                    thread_id=thread.id,
                    role="user",
                    content=f"Here is the CSV data for {selected_month}:\n\n{csv_data}\n\nPlease run the monthly marketing analysis."
                )

                run = openai.beta.threads.runs.create(
                    thread_id=thread.id,
                    assistant_id=assistant_id
                )

                while run.status not in ["completed", "failed", "cancelled"]:
                    time.sleep(1)
                    run = openai.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)

                messages = openai.beta.threads.messages.list(thread_id=thread.id)
                result = messages.data[0].content[0].text.value

                st.session_state.analysis_output = result
                st.session_state.conversation_history = [
                    {"role": "user", "content": f"(CSV data submitted for {selected_month})"},
                    {"role": "assistant", "content": result}
                ]



    # === Show Output, Follow-ups, Reset, and Export (only after analysis)
    if st.session_state.analysis_output:
        if st.session_state.get("rerun_requested"):
            del st.session_state.rerun_requested
            st.experimental_rerun()
        st.subheader(f"🧠 GPT Analysis for {selected_month}")
        st.markdown(st.session_state.analysis_output)

        # === Follow-up Questions
        if st.session_state.followup_count < 3:
            followup = st.text_input("Ask a follow-up question about the data", key=f"followup_input_{st.session_state.followup_count}")
            if followup:
                with st.spinner("Thinking..."):
                    preview_csv = "\n".join(st.session_state.csv_data.split("\n")[:11])
                    followup_prompt = f"""Reminder: this follow-up relates to the CSV data for {selected_month}:\n\n{preview_csv}\n\nQuestion: {followup}"""
                    st.session_state.conversation_history.append({"role": "user", "content": followup_prompt})

                    openai.beta.threads.messages.create(
                        thread_id=st.session_state.analysis_thread_id,
                        role="user",
                        content=followup_prompt
                    )

                    followup_run = openai.beta.threads.runs.create(
                        thread_id=st.session_state.analysis_thread_id,
                        assistant_id=assistant_id
                    )

                    while followup_run.status not in ["completed", "failed", "cancelled"]:
                        time.sleep(1)
                        followup_run = openai.beta.threads.runs.retrieve(
                            thread_id=st.session_state.analysis_thread_id,
                            run_id=followup_run.id
                        )

                    followup_messages = openai.beta.threads.messages.list(thread_id=st.session_state.analysis_thread_id)
                    reply = followup_messages.data[0].content[0].text.value

                    st.session_state.conversation_history.append({"role": "assistant", "content": reply})
                    st.session_state.followup_count += 1

                    st.session_state.last_reply = f"#### 💬 GPT Reply #{st.session_state.followup_count}\n\n{reply}"
                    st.session_state.rerun_requested = True
                    st.stop()

        elif st.session_state.followup_count >= 3:
            st.info("🔒 Follow-up questions limit reached. Please click '🔁 Reset Analysis' to start over.")

        if "last_reply" in st.session_state:
            st.markdown(st.session_state.last_reply)

        # === Reset Button
        if st.button("🔁 Reset Analysis"):
            st.session_state.analysis_output = None
            st.session_state.followup_count = 0
            st.session_state.conversation_history = []
            st.session_state.last_reply = None
            st.session_state.csv_data = None
            st.rerun()

        # === Export to .docx
        if st.button("📄 Export to .docx"):
            doc = Document()
            doc.add_heading(f"Marketing Analysis – {selected_month}", 0)
            for msg in st.session_state.conversation_history:
                role = msg["role"]
                if role == "user":
                    doc.add_paragraph(f"User:\n{msg['content']}")
                elif role == "assistant":
                    doc.add_paragraph(f"Kimbell Analysis:\n{msg['content']}")
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            filename = f"Analysis_{selected_month}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.download_button("📄 Download .docx File", data=buffer, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")